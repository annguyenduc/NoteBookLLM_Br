
import os
import sys
from pathlib import Path
import markdown
import re
from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
ROOT = Path("d:/NoteBookLLM_Br")
CSS_PATH = ROOT / "scratch/k10_premium_style.css"
OUTPUT_DIR = ROOT / "distilled/K10_Pedagogical_Kit"

# Subjects and Files Mapping
MATERIALS = {
    "Math_10": [
        ROOT / "3-resources/wiki/ATOMS_Prompt_Engineering_K10_Toan.md",
        ROOT / "3-resources/process/Math10_Prompt_Comparison_Examples.md",
        ROOT / "3-resources/process/Math10_Prompt_Showcase.md",
        ROOT / "3-resources/process/Lesson_Plan_Math10_Prompt_45m.md",
        ROOT / "3-resources/process/Math10_Gemini_Study_Guide.md"
    ],
    "English_10": [
        ROOT / "3-resources/wiki/ATOMS_Prompt_Engineering_K10_Anh.md",
        ROOT / "3-resources/process/English10_Prompt_Comparison_Examples.md",
        ROOT / "3-resources/process/English10_Prompt_Showcase.md",
        ROOT / "3-resources/process/Lesson_Plan_English10_Prompt_45m.md"
    ],
    "Literature_10": [
        ROOT / "3-resources/wiki/ATOMS_Prompt_Engineering_K10_Van.md",
        ROOT / "3-resources/process/Literature10_Prompt_Comparison_Examples.md",
        ROOT / "3-resources/process/Literature10_Prompt_Showcase.md",
        ROOT / "3-resources/process/Lesson_Plan_Literature10_Prompt_45m.md"
    ]
}

COPY_SCRIPT = """
<script>
function copyPrompt(text, btn) {
    // Xóa phần nguồn khi copy để học sinh chỉ dán Prompt thuần túy
    const cleanText = text.split('📖 Nguồn:')[0].replace(/"/g, '').trim();
    navigator.clipboard.writeText(cleanText).then(() => {
        const originalText = btn.innerHTML;
        btn.classList.add('success');
        btn.innerHTML = '✅ ĐÃ SAO CHÉP!';
        btn.style.backgroundColor = '#2e7d32';
        setTimeout(() => {
            btn.classList.remove('success');
            btn.innerHTML = originalText;
            btn.style.backgroundColor = '#1a237e';
        }, 2000);
    });
}

document.addEventListener('DOMContentLoaded', () => {
    // Xử lý các Prompt nằm trong Bảng (Tables)
    const cells = document.querySelectorAll('td');
    cells.forEach((cell) => {
        const text = cell.innerText;
        if (text.includes('[Role]') || text.includes('[Task]')) {
            const btn = document.createElement('button');
            btn.innerHTML = '📋 SAO CHÉP PROMPT';
            btn.className = 'copy-btn-inline';
            btn.style = 'display: block; margin-top: 15px; padding: 12px; cursor: pointer; background: #1a237e; color: white; border: none; border-radius: 8px; font-weight: bold; width: 100%; box-shadow: 0 4px 6px rgba(0,0,0,0.1); box-sizing: border-box;';
            btn.onclick = () => copyPrompt(text, btn);
            cell.appendChild(btn);
        }
    });

    // Xử lý các Prompt nằm trong hộp thoại (Code blocks/Pre) nếu có
    const codeBlocks = document.querySelectorAll('pre code');
    codeBlocks.forEach((code) => {
        const text = code.innerText;
        if (text.includes('[Role]') || text.includes('[Task]')) {
            const btn = document.createElement('button');
            btn.innerHTML = '📋 SAO CHÉP PROMPT';
            btn.className = 'copy-btn-inline';
            btn.style = 'display: block; margin: 10px 0; padding: 12px; cursor: pointer; background: #1a237e; color: white; border: none; border-radius: 8px; font-weight: bold; width: 100%; box-shadow: 0 4px 6px rgba(0,0,0,0.1); box-sizing: border-box;';
            btn.onclick = () => copyPrompt(text, btn);
            code.parentElement.parentElement.insertBefore(btn, code.parentElement);
        }
    });
});
</script>
"""

MATHJAX_SCRIPT = """
<script>
window.MathJax = {
  tex: {
    inlineMath: [['$', '$'], ['\\\\(', '\\\\)']],
    displayMath: [['$$', '$$'], ['\\\\[', '\\\\]']]
  },
  svg: { fontCache: 'global' }
};
</script>
<script id="MathJax-script" async src="https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js"></script>
"""

def load_css():
    with open(CSS_PATH, "r", encoding="utf-8") as f:
        return f.read()

def clean_latex_for_docx(text):
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    text = re.sub(r'\\frac\{(.*?)\}\{(.*?)\}', r'\1/\2', text)
    text = re.sub(r'\\text\{(.*?)\}', r'\1', text)
    text = text.replace('^2', '²').replace('^3', '³').replace('^n', 'ⁿ')
    text = text.replace('\\cdot', '·').replace('\\times', '×').replace('\\le', '≤').replace('\\ge', '≥').replace('\\vec', 'Véc-tơ ').replace('\\Delta', 'Δ')
    return text

def md_to_docx(md_content, output_path):
    doc = Document()
    lines = md_content.split('\n')
    for line in lines:
        line = clean_latex_for_docx(line)
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.strip():
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(output_path)

def generate_index_dashboard(css):
    print("Generating Master Dashboard (index.html)...")
    content = f"""
    <!DOCTYPE html>
    <html lang="vi">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Gemini Express K10 Hub</title>
        <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@700&family=Inter:wght@400;600&display=swap" rel="stylesheet">
        <style>
            {css}
            body {{ font-family: 'Inter', sans-serif; background: #f4f7f9; margin: 0; padding: 0; }}
            .hero {{ background: #1a237e; color: white; padding: 40px 20px; text-align: center; }}
            .hero h1 {{ font-family: 'Montserrat', sans-serif; margin: 0; font-size: 2rem; }}
            .container {{ max-width: 900px; margin: 20px auto; padding: 20px; }}
            .grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap: 20px; }}
            .card {{ background: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); transition: transform 0.2s; text-decoration: none; color: inherit; border-top: 5px solid #1a237e; }}
            .card:hover {{ transform: translateY(-5px); }}
            .card h2 {{ margin: 0 0 10px 0; color: #1a237e; font-size: 1.25rem; }}
            .card p {{ font-size: 0.9rem; color: #666; margin-bottom: 15px; }}
            .footer {{ text-align: center; padding: 20px; color: #aaa; font-size: 0.8rem; }}
        </style>
    </head>
    <body>
        <div class="hero">
            <h1>🚀 Gemini Express K10 Hub</h1>
            <p>Trạm học liệu thông minh dành cho học sinh lớp 10</p>
        </div>
        <div class="container">
            <div class="grid">
                <a href="distilled/K10_Pedagogical_Kit/Math_10/Math10_Gemini_Study_Guide.html" class="card">
                    <h2>📐 Toán Học 10</h2>
                    <p>Ôn tập chương trình HKII, hình học tọa độ, tổ hợp và xác suất.</p>
                </a>
                <a href="distilled/K10_Pedagogical_Kit/English_10/English10_Prompt_Showcase.html" class="card">
                    <h2>🇬🇧 Tiếng Anh 10</h2>
                    <p>Phát triển kỹ năng Writing, Speaking và từ vựng theo chủ điểm.</p>
                </a>
                <a href="distilled/K10_Pedagogical_Kit/Literature_10/Literature10_Prompt_Showcase.html" class="card">
                    <h2>📚 Ngữ Văn 10</h2>
                    <p>Phân tích tác phẩm, viết văn nghị luận và tư duy sáng tạo.</p>
                </a>
            </div>
        </div>
        <div class="footer">Dự án Pedagogical AI Toolkit © 2026</div>
    </body>
    </html>
    """
    with open(ROOT / "index.html", "w", encoding="utf-8") as f:
        f.write(content)

def export():
    css = load_css()
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    generate_index_dashboard(css)
    
    with sync_playwright() as p:
        browser = p.chromium.launch()
        
        for subject, files in MATERIALS.items():
            sub_dir = OUTPUT_DIR / subject
            os.makedirs(sub_dir, exist_ok=True)
            
            for file_path in files:
                if not file_path.exists(): continue
                base_name = file_path.stem
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                
                # HTML
                html_body = markdown.markdown(content, extensions=['tables', 'fenced_code'])
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <meta name="viewport" content="width=device-width, initial-scale=1.0">
                    <title>{base_name}</title>
                    {MATHJAX_SCRIPT}
                    {COPY_SCRIPT}
                    <style>{css}</style>
                </head>
                <body>
                    <div class="mobile-nav" style="background:#1a237e; color:white; padding:10px; position:sticky; top:0; text-align:center; z-index:1000;">
                        <a href="../../../index.html" style="color:white; text-decoration:none; font-weight:bold;">← Quay về Dashboard</a>
                    </div>
                    <div style="padding: 20px;">
                        {html_body}
                    </div>
                </body>
                </html>
                """
                
                html_out = sub_dir / f"{base_name}.html"
                with open(html_out, "w", encoding="utf-8") as hf:
                    hf.write(full_html)
                
                # PDF
                pdf_out = sub_dir / f"{base_name}.pdf"
                page = browser.new_page()
                page.set_content(full_html)
                page.wait_for_timeout(3000) 
                page.pdf(path=str(pdf_out), format="A4", print_background=True, 
                         margin={"top": "15mm", "bottom": "15mm", "left": "15mm", "right": "15mm"})
                page.close()
                
                # DOCX
                docx_out = sub_dir / f"{base_name}.docx"
                md_to_docx(content, docx_out)

        browser.close()

if __name__ == "__main__":
    export()
    print("Mass Export Complete with Gemini Express Hub Integration!")
